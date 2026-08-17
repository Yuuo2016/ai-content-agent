"""
题二：内容运营 Agent（Content Operations Agent）

完整工作流（7 步）：
    自动发现热门话题 → 抓取真实信息源 → AI 基于热门话题+真实信息生成选题
    → 逐条选题审核 → AI 生成多平台内容 → 生成发布计划
    → 发布前人工确认 → 多渠道推送（飞书 + QQ 邮箱）+ 生成 Word 文档

覆盖题目要求：
    - 使用真实信息来源（抓取 RSS/GitHub 等真实资讯，非纯 AI 编造）
    - 收集 + 生成（4 个任务中选 2 个）
    - 对外发布前人工确认（发布前确认节点）
并实现加分项：
    - 自动找热门话题（知乎热榜/微博热搜/百度热搜/V2Ex/HN）
    - 生成平台化版本（小红书/公众号风格）
    - 生成发布计划/报告

运行方式：
    python problem2_content/main.py
"""
import os
import sys
import json
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from common import llm, push
from problem2_content.review_topics import review_topics
from problem2_content.publish_confirm import publish_confirm
from problem2_content.sources import fetch_all, fetch_hot_topics

def collect_topics(source_items: list, hot_topics: list = None, keyword: str = "", count: int = 5) -> list:
    """AI 基于热门话题 + 真实抓取信息生成选题。

    把热门话题榜和真实资讯一起作为输入，让 AI 从中提炼高热度选题，
    每个选题都带来源(source)与链接(url)，保证「真实信息来源」可追溯。
    """
    hot_topics = hot_topics or []
    candidates = [
        {"title": it["title"], "source": it["source"], "summary": it["summary"], "url": it["url"]}
        for it in source_items
    ]
    hot_list = [
        {"title": ht["title"], "source": ht["source"], "url": ht.get("url", "")}
        for ht in hot_topics
    ]

    messages = [
        {
            "role": "system",
            "content": (
                "你是海外 AI 产品的内容运营专家。请基于下方提供的热门话题榜和真实资讯来源，"
                "提炼出适合「AI 产品出海增长、科技财经内容」的高热度选题。"
                "优先结合热门话题与资讯的交叉点（即既在热搜上、又有真实报道的话题）。"
                "每个选题包含标题(title)、目标人群(audience)、内容角度(angle)、"
                "平台建议(platform)、选题理由(rationale)，并原样保留其来源(source)与"
                "链接(url)。只输出 JSON 数组。"
            ),
        },
        {
            "role": "user",
            "content": (
                f"请围绕主题「{keyword or 'AI 产品出海增长'}」从以下信息中提炼 "
                f"{count} 个热门内容选题，要求贴合海外市场（欧美/东南亚/日韩）用户关注点，"
                "每个选题给出针对性理由，并标注其来源与链接：\n\n"
                f"--- 热门话题榜 ---\n{json.dumps(hot_list, ensure_ascii=False, indent=2)}\n\n"
                f"--- 真实资讯 ---\n{json.dumps(candidates, ensure_ascii=False, indent=2)}"
            ),
        },
    ]
    try:
        result = llm.chat_json(messages, temperature=0.7, max_tokens=2000)
        if isinstance(result, dict):
            result = result.get("topics") or result.get("items") or []
        return result[:count]
    except Exception as e:
        print(f"[错误] 选题生成失败: {e}")
        # 兜底：优先用热门话题，其次用资讯
        fallback = hot_topics[:count] if hot_topics else source_items[:count]
        return [
            {
                "title": f"{keyword or 'AI出海'} 选题{i+1}",
                "audience": "海外用户",
                "angle": "行业趋势",
                "platform": "公众号",
                "rationale": f"基于热门话题「{it['title'][:30] or '行业动态'}」提炼",
                "source": it.get("source", ""),
                "url": it.get("url", ""),
            }
            for i, it in enumerate(fallback)
        ]

def generate_content(topic: dict, platform: str = "公众号") -> str:
    """AI 根据选题生成平台化内容"""
    messages = [
        {
            "role": "system",
            "content": (
                "你是资深内容运营，擅长为不同平台创作高转化内容。"
                "请根据选题和平台要求，输出完整可发布的内容正文，"
                "包括吸引人的开头、结构化正文、结尾行动号召。"
            ),
        },
        {
            "role": "user",
            "content": (
                f"选题标题：{topic.get('title', '')}\n"
                f"目标人群：{topic.get('audience', '')}\n"
                f"内容角度：{topic.get('angle', '')}\n"
                f"信息来源：{topic.get('source', '')}\n"
                f"来源链接：{topic.get('url', '')}\n"
                f"目标平台：{platform}\n"
                f"请为「{platform}」平台创作一篇完整内容（{platform}风格，"
                f"注意标题党适度、段落短小、可读性强，并在文中适当引用来源）。"
            ),
        },
    ]
    return llm.chat(messages, temperature=0.7, max_tokens=1500)

def generate_multi_platform(topic: dict, platforms: list = None) -> dict:
    """加分项：为选题生成多个平台的版本。

    Returns:
        {platform: content}
    """
    platforms = platforms or ["公众号", "小红书"]
    versions = {}
    for p in platforms:
        print(f"      正在生成「{p}」版本...")
        try:
            versions[p] = generate_content(topic, platform=p)
        except Exception as e:
            print(f"      [错误]「{p}」版本生成失败: {e}")
            versions[p] = f"（{p}版本生成失败）"
    return versions

def generate_publish_plan(topics: list, versions_map: dict) -> str:
    """加分项：生成发布计划/报告文本。

    Args:
        topics: 已审核通过的选题列表
        versions_map: {topic_title: {platform: content}}

    Returns:
        发布计划文本
    """
    lines = [f"📝 内容发布计划（{datetime.now().strftime('%Y-%m-%d')}）", ""]
    for i, t in enumerate(topics, 1):
        lines.append(f"选题{i}: {t.get('title', '')}")
        lines.append(f"   目标人群: {t.get('audience', '')}")
        lines.append(f"   内容角度: {t.get('angle', '')}")
        if t.get("source"):
            lines.append(f"   信息来源: {t.get('source', '')}")
        if t.get("url"):
            lines.append(f"   来源链接: {t.get('url', '')}")
        versions = versions_map.get(t.get("title", ""), {})
        if versions:
            lines.append(f"   生成平台版本: {', '.join(versions.keys())}")
            lines.append(f"   建议发布时间: 工作日 12:00 / 20:00（按平台活跃时段）")
            lines.append(f"   发布节奏: 首选平台发布后 2 小时转载到次平台")
        lines.append("")
    return "\n".join(lines)

def save_to_word(title: str, content: str, path: str) -> str:
    """把内容保存为 Word 文档"""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    for para in content.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.save(path)
    return path

def main():
    print("=" * 60)
    print("题二：内容运营 Agent 启动")
    print("=" * 60)

    # 1. 加分项：自动发现热门话题
    print("\n[1/7] 自动发现热门话题...")
    print("  正在抓取知乎热榜 / 微博热搜 / 百度热搜 / V2Ex / HN ...")
    hot_topics = fetch_hot_topics()
    print(f"  发现 {len(hot_topics)} 个热门话题：")
    for i, ht in enumerate(hot_topics[:15], 1):
        print(f"    {i}. [{ht['source']}] {ht['title'][:50]}")

    # 2. 抓取真实信息源
    print("\n[2/7] 抓取真实信息源...")
    source_items = fetch_all(max_items=20)
    print(f"      共抓取到 {len(source_items)} 条真实资讯")
    for it in source_items[:5]:
        print(f"      - [{it['source']}] {it['title'][:40]}")

    # 3. AI 基于热门话题 + 真实信息生成选题
    print("\n[3/7] AI 基于热门话题 + 真实信息生成选题...")
    keyword = input("请输入内容主题（直接回车使用默认「AI 产品出海增长」）: ").strip()
    topics = collect_topics(source_items, hot_topics, keyword, count=5)
    print(f"      生成 {len(topics)} 个选题（均带真实来源）")

    # 4. 逐条选题审核（e/r 不推送，p 才进入下一步）
    print("\n[4/7] 进入选题审核...")
    final_topics = review_topics(topics)
    if not final_topics:
        print("所有选题均被拒绝，流程结束。")
        return

    # 5. 为每个选题生成多平台内容
    print("\n[5/7] AI 生成多平台内容...")
    versions_map = {}
    for t in final_topics:
        print(f"  - 正在为「{t.get('title')}」生成内容...")
        versions_map[t.get("title", "")] = generate_multi_platform(t)

    # 6. 生成发布计划
    print("\n[6/7] 生成发布计划...")
    plan = generate_publish_plan(final_topics, versions_map)
    print(plan)

    # 7. 输出：发布前确认 → 多渠道推送（飞书 + QQ 邮箱）+ 生成 Word 文档
    print("\n[7/7] 输出内容...")
    os.makedirs("output", exist_ok=True)

    report = plan + "\n===== 内容正文 =====\n"
    for t in final_topics:
        report += f"\n【{t.get('title')}】\n"
        for p, content in versions_map.get(t.get("title", ""), {}).items():
            report += f"\n-- {p} 版本 --\n{content}\n"

    # 7a. 发布前人工确认（硬性要求：对外发布前需人工确认）
    confirmed = publish_confirm("内容运营报告", report)
    if not confirmed:
        print(">>> 已拒绝发布，流程结束（未推送）。")
        return

    # 7b. 确认后：多渠道推送（飞书 + QQ 邮箱，企业微信默认不发）
    try:
        results = push.push_all("内容运营报告", confirmed)
        for channel, res in results.items():
            # 题二不显示企业微信（即使 push_all 返回了也跳过）
            if channel == "企业微信":
                continue
            status = "✅ 成功" if res.get("ok") else "❌ 失败/跳过"
            print(f"      [{channel}] {status}  {res.get('detail')}")
    except Exception as e:
        print(f"      [错误] 多渠道推送失败: {e}")

    # 7c. 生成 Word 文档
    doc_path = f"output/content_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    try:
        save_to_word("内容运营报告", confirmed, doc_path)
        print(f"      ✅ Word 文档已生成: {doc_path}")
    except Exception as e:
        print(f"      [错误] Word 生成失败: {e}")

    print("\n流程完成。")

if __name__ == "__main__":
    main()